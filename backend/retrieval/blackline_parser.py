import docx
from typing import List, Dict

class RedlineParser:
    """
    Implements logic from the Mike-redline fork. 
    Standard LLM document ingestion flattens text and loses tracked changes.
    This parser extracts <w:ins> and <w:del> tags from DOCX XML to make the LLM
    "redline-aware" for contract negotiation workflows.
    """
    
    def parse_docx_redlines(self, file_path: str) -> Dict[str, List[str]]:
        """
        Parses a DOCX file and specifically extracts insertions and deletions.
        """
        try:
            doc = docx.Document(file_path)
            insertions = []
            deletions = []
            
            # Note: A true deep parse requires diving into the lxml elements of python-docx.
            # This is a high-level representation of the logic:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # In python-docx, tracked changes are deeper in the XML tree (run._r)
                    if '<w:ins' in run._r.xml:
                        insertions.append(run.text)
                    elif '<w:delText' in run._r.xml:
                        deletions.append(run.text)
                        
            return {
                "insertions": insertions,
                "deletions": deletions,
                "full_text": "\\n".join([p.text for p in doc.paragraphs])
            }
        except Exception as e:
            return {"error": str(e)}
